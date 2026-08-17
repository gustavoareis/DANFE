import random
import re
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt
import holidays
from num2words import num2words
import pdfplumber

# --- Constantes ---
MESES = [
    "",
    "JANEIRO",
    "FEVEREIRO",
    "MARÇO",
    "ABRIL",
    "MAIO",
    "JUNHO",
    "JULHO",
    "AGOSTO",
    "SETEMBRO",
    "OUTUBRO",
    "NOVEMBRO",
    "DEZEMBRO",
]


# --- Helpers Utilitários ---
def parse_float(val) -> float:
  """Converte string formatada em BRL (ex: '1.234,56') para float."""
  clean = re.sub(r"[^\d,.-]", "", str(val or "")).replace(".", "").replace(",", ".")
  return float(clean) if clean else 0.0


def format_brl(val: float) -> str:
  return f"{val:,.2f}".translate(str.maketrans(",.", ".,"))


def format_extenso(val: float) -> str:
  return num2words(val, lang="pt_BR", to="currency").upper()


def retroceder_dias_uteis(data_ref: datetime, dias: int = 3) -> datetime:
  """Retrocede dias úteis considerando feriados estaduais do Ceará."""
  feriados = holidays.BR(subdiv="CE", years=data_ref.year)
  feriados.update({
      datetime(data_ref.year, 4, 13).date(): "Aniversário Fortaleza",
      datetime(data_ref.year, 8, 15).date(): "Nossa Senhora",
  })

  atual = data_ref
  while dias > 0:
    atual -= timedelta(days=1)
    if atual.weekday() < 5 and atual.date() not in feriados:
      dias -= 1
  return atual


# --- Processamento de PDF ---


def extrair_dados_pdf_produto(caminho_pdf: Path, txt_extraido: str) -> dict:
  """Lógica original para notas fiscais de produtos (NF-e)."""
  dados = {
      "numero_nf": "000",
      "cliente": "",
      "cnpj": "",
      "data_emissao": datetime.now(),
      "itens": [],
      "obs": "",
  }
  
  # Extração de tabelas
  with pdfplumber.open(caminho_pdf) as pdf:
    for page in pdf.pages:
      for table in page.extract_tables():
        for row in table:
          if row and len(row) >= 9 and str(row[0]).isdigit():
            try:
              dados["itens"].append({
                  "descricao": str(row[1] or "").replace("\n", " "),
                  "und": str(row[5] or ""),
                  "qtd": parse_float(row[6]),
                  "vlr_unit": parse_float(row[7]),
                  "vlr_total": parse_float(row[8]),
              })
            except (ValueError, TypeError):
              continue

  # Extração de cabeçalho via Regex
  if m := re.search(r"Nº\s*(\d+)", txt_extraido):
    dados["numero_nf"] = m.group(1)
  if m := re.search(r"DATA DE EMISSÃO\n*(\d{2}/\d{2}/\d{4})|(\d{2}/\d{2}/\d{4})", txt_extraido):
    dados["data_emissao"] = datetime.strptime(m.group(1) or m.group(2), "%d/%m/%Y")
  if m := re.search(r"CNPJ/CPF[^\d]*(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", txt_extraido):
    dados["cnpj"] = m.group(1).strip()
  if m := re.search(r"OBS:\s*([^\n]+)", txt_extraido, re.IGNORECASE):
    dados["obs"] = m.group(1).strip()

  linhas = txt_extraido.split("\n")
  for i, linha in enumerate(linhas):
    if "NOME/RAZÃO SOCIAL" in linha:
      resto = linha.split("NOME/RAZÃO SOCIAL")[-1].strip()
      dados["cliente"] = (
          resto if (resto and "CNPJ" not in resto)
          else (re.sub(r"\s*\d{2}\.\d{3}\.\d{3}/.*", "", linhas[i + 1]).strip() if i + 1 < len(linhas) else "")
      )
      break

  return dados

def extrair_dados_pdf_servico(txt_extraido: str) -> dict:
  """Nova lógica dedicada à Nota Fiscal de Serviço (NFS-e) de Fortaleza."""
  dados = {
      "numero_nf": "000",
      "cliente": "",
      "cnpj": "",
      "data_emissao": datetime.now(),
      "itens": [],
      "obs": "",
  }

  # 1. Número da NFS-e
  if m := re.search(r"Número da\s*NFS-e\s*(\d+)", txt_extraido, re.IGNORECASE):
      dados["numero_nf"] = m.group(1)
  elif m := re.search(r"NFS-e\s*(\d+)", txt_extraido):
      dados["numero_nf"] = m.group(1)

  # 2. Data de Emissão
  if m := re.search(r"(\d{2}/\d{2}/\d{4})", txt_extraido):
      dados["data_emissao"] = datetime.strptime(m.group(1), "%d/%m/%Y")

  # 3. CNPJ do Tomador
  if m := re.search(r"(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", txt_extraido):
      dados["cnpj"] = m.group(1)

  # 4. Cliente e Observação a partir da discriminação
  descricao_bruta = ""
  if m := re.search(r"DISCRIMINAÇÃO DOS SERVIÇOS\n(.*?)(?=\n(-)?[A-Z]|\nValor|\nCálculo)", txt_extraido, re.DOTALL | re.IGNORECASE):
      descricao_bruta = m.group(1).replace("\n", " ").strip()
  else:
      if m := re.search(r"DISCRIMINAÇÃO DOS SERVIÇOS(.*?)Valor", txt_extraido, re.DOTALL | re.IGNORECASE):
           descricao_bruta = m.group(1).replace("\n", " ").strip()

  if "OBS:" in descricao_bruta.upper():
      partes = re.split(r"OBS:", descricao_bruta, flags=re.IGNORECASE)
      descricao_limpa = partes[0].strip()
      dados["obs"] = partes[1].strip()
  else:
      descricao_limpa = descricao_bruta

  # CORREÇÃO 1: Remove textos dinâmicos como "NO VALOR DE R$ 682,80" da descrição
  descricao_limpa = re.sub(r"\s*NO VALOR DE R\$\s*[\d.,]+", "", descricao_limpa, flags=re.IGNORECASE).strip()

  # 5. Nome do Cliente (Tomador)
  if m := re.search(r"DADOS DO TOMADOR DE SERVIÇOS[^\n]*\n([^\n]+)\n", txt_extraido):
      linhas = txt_extraido.split("\n")
      for i, linha in enumerate(linhas):
          if "DADOS DO TOMADOR DE SERVIÇOS" in linha:
              for possivel_cliente in linhas[i+1:i+5]:
                  if possivel_cliente and len(possivel_cliente) > 5 and "CNPJ" not in possivel_cliente:
                      # CORREÇÃO 2: Remove o prefixo "Razão Social/Nome" se o PDF aglutinar na mesma linha
                      cliente_limpo = re.sub(r"^Razão Social/Nome\s*", "", possivel_cliente.strip(), flags=re.IGNORECASE)
                      dados["cliente"] = cliente_limpo.strip()
                      break
              break

  # 6. Valor Líquido ou Valor do Serviço
  vlr_total = 0.0
  if m := re.search(r"Valor Líquido R\$\s*([\d.,]+)", txt_extraido):
       vlr_total = parse_float(m.group(1))
  elif m := re.search(r"Valor dos Serviços R\$\s*([\d.,]+)", txt_extraido):
       vlr_total = parse_float(m.group(1))

  # 7. Registra o serviço na tabela (Qtde: 1)
  dados["itens"].append({
      "descricao": descricao_limpa if descricao_limpa else "PRESTAÇÃO DE SERVIÇOS",
      "und": "SV",
      "qtd": 1.0,
      "vlr_unit": vlr_total,
      "vlr_total": vlr_total
  })

  return dados

def extrair_dados_pdf(caminho_pdf: Path) -> dict:
  """Função roteadora: extrai o texto base e direciona para a função correta."""
  txt_extraido = ""
  with pdfplumber.open(caminho_pdf) as pdf:
      for page in pdf.pages:
          txt_extraido += (page.extract_text() or "") + "\n"
          
  is_servico = "NOTA FISCAL DE SERVIÇO" in txt_extraido.upper() or "NFS-e" in txt_extraido
  
  if is_servico:
      return extrair_dados_pdf_servico(txt_extraido)
  else:
      return extrair_dados_pdf_produto(caminho_pdf, txt_extraido)
  

def reajustar_valores(itens: list, percentual: float) -> tuple[list, float]:
  """Aplica margem com reajuste no último item para arredondamento exato."""
  total_alvo = round(sum(i["vlr_total"] for i in itens) * (1 + percentual), 2)
  novos_itens, soma = [], 0.0

  for i in itens:
    fator = random.uniform(
        1.0 + max(0.005, percentual - 0.02), 1.0 + percentual + 0.02
    )
    unit = round(i["vlr_unit"] * fator, 2)
    total = round(unit * i["qtd"], 2)
    novos_itens.append({**i, "vlr_unit": unit, "vlr_total": total})
    soma += total

  if novos_itens and (dif := round(total_alvo - soma, 2)) != 0:
    novos_itens[-1]["vlr_total"] = round(
        novos_itens[-1]["vlr_total"] + dif, 2
    )
    if novos_itens[-1]["qtd"] > 0:
      novos_itens[-1]["vlr_unit"] = round(
          novos_itens[-1]["vlr_total"] / novos_itens[-1]["qtd"], 2
      )

  return novos_itens, total_alvo


# --- Manipulação de Word ---
def substituir_texto(paragrafo, novo_texto: str):
  """Substitui o texto preservando a formatação do primeiro run."""
  if not paragrafo.runs:
    paragrafo.text = novo_texto
    return
  paragrafo.runs[0].text = novo_texto
  for r in paragrafo.runs[1:]:
    r.text = ""


def style_table(tabela, is_jw: bool):
  """Aplica Mongolian Baiti 9pt em toda a tabela. Se for JW, define TODA a tabela em negrito."""
  for idx_row, row in enumerate(tabela.rows):
    is_header = idx_row == 0
    for cell in row.cells:
      for p in cell.paragraphs:
        for r in p.runs:
          r.font.name = "Mongolian Baiti"
          r.font.size = Pt(9)
          r.font.bold = True if is_jw else is_header


def preencher_tabela(tabela, itens: list, is_jw: bool):
  """Preenche os dados da tabela, mantendo alinhamentos e aplicando estilo."""
  if not tabela.rows:
    return

  headers = [c.text.upper() for c in tabela.rows[0].cells]
  col_map = {
      "num": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["ORDEM", "ITEM", "Nº"])
          ),
          0,
      ),
      "desc": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["DESC", "PRODUTO", "SERVIÇO"])
          ),
          1,
      ),
      "und": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["UND", "UNID"])
          ),
          2,
      ),
      "qtd": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["QTD", "QUANT"])
          ),
          3,
      ),
      "unit": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["UNIT", "PR. UNIT"])
          ),
          4,
      ),
      "total": next(
          (
              i
              for i, h in enumerate(headers)
              if any(x in h for x in ["TOTAL", "PR. TOTAL"])
          ),
          5,
      ),
  }

  alinhamentos = {
      col_idx: tabela.rows[1].cells[col_idx].paragraphs[0].alignment
      for col_idx in col_map.values()
      if len(tabela.rows) > 1 and tabela.rows[1].cells[col_idx].paragraphs
  }

  for idx, item in enumerate(itens):
    if idx + 1 >= len(tabela.rows):
      tabela.add_row()
    cells = tabela.rows[idx + 1].cells

    valores = {
        "num": str(idx + 1),
        "desc": item["descricao"],
        "und": item["und"],
        "qtd": (
            f"{item['qtd']:.0f}"
            if item["qtd"].is_integer()
            else f"{item['qtd']:.2f}".replace(".", ",")
        ),
        "unit": f"R$ {format_brl(item['vlr_unit'])}",
        "total": f"R$ {format_brl(item['vlr_total'])}",
    }

    for chave, col_idx in col_map.items():
      if col_idx < len(cells):
        cells[col_idx].text = valores[chave]
        if align := alinhamentos.get(col_idx):
          if cells[col_idx].paragraphs:
            cells[col_idx].paragraphs[0].alignment = align

  style_table(tabela, is_jw)


def preencher_word(
    template_path: Path,
    output_path: Path,
    dados: dict,
    itens: list,
    total: float,
    is_jw: bool,
):
  """Gera a proposta preenchendo os dados no Word."""
  doc = Document(template_path)
  dt = retroceder_dias_uteis(dados["data_emissao"], 3)
  data_fmt = f"{dt.strftime('%d')} DE {MESES[dt.month]} DE {dt.year}"
  val_str, ext_str = format_brl(total), format_extenso(total)

  for p in doc.paragraphs:
    txt = p.text.upper()
    if txt.startswith("FORTALEZA") and " DE " in txt:
      substituir_texto(
          p, f"{'FORTALEZA' if p.text.isupper() else 'Fortaleza'}, {data_fmt}"
      )
    elif "CLIENTE:" in txt and dados.get("cliente"):
      partes = [f"Cliente: {dados['cliente']}"] + (
          [dados["obs"]] if dados.get("obs") else []
      )
      if is_jw and dados.get("cnpj"):
        partes.append(f"CNPJ: {dados['cnpj']}")
      substituir_texto(p, "\n".join(partes))
    elif "CNPJ:" in txt and not is_jw and dados.get("cnpj"):
      substituir_texto(p, f"CNPJ: {dados['cnpj']}")
    elif is_jw and "VALOR DA PROPOSTA" in txt:
      substituir_texto(p, f"Valor da Proposta R$ {val_str} ({ext_str})")
    elif not is_jw and "VALOR TOTAL DA PROPOSTA:" in txt:
      substituir_texto(p, f"VALOR TOTAL DA PROPOSTA:  R$ {val_str}")

  if doc.tables:
    preencher_tabela(doc.tables[0], itens, is_jw)

  doc.save(output_path)


# --- Fluxo Principal ---
def processar_notas():
  entrada, saida = Path("notas_entrada"), Path("propostas_saida")
  tpl_lc = (
      Path("template_lc.docx")
      if Path("template_lc.docx").exists()
      else Path("template.docx")
  )
  tpl_jw = Path("template_jw.docx")

  if not entrada.exists():
    entrada.mkdir()
    return print(
        f"Criando pasta '{entrada}'. Adicione os PDFs nela e execute novamente."
    )

  saida.mkdir(exist_ok=True)
  arquivos = list(entrada.glob("*.pdf"))

  if not arquivos:
    return print("Nenhum PDF encontrado em 'notas_entrada'.")

  for arquivo in arquivos:
    print(f"\n--- Processando: {arquivo.name} ---")
    dados = extrair_dados_pdf(arquivo)
    pasta_nf = saida / f"NF {dados['numero_nf']}"
    pasta_nf.mkdir(exist_ok=True)

    if tpl_lc.exists():
      itens_lc, tot_lc = reajustar_valores(dados["itens"], 0.03)
      preencher_word(
          tpl_lc,
          pasta_nf / f"LC COMERCIAL NF {dados['numero_nf']}.docx",
          dados,
          itens_lc,
          tot_lc,
          is_jw=False,
      )
      print(
          f"  ✓ LC criada (+3.0%) - Data:"
          f" {dados['data_emissao'].strftime('%d/%m/%Y')}"
      )

    if tpl_jw.exists():
      margem_jw = random.uniform(0.05, 0.08)
      itens_jw, tot_jw = reajustar_valores(dados["itens"], margem_jw)
      preencher_word(
          tpl_jw,
          pasta_nf / f"JW COMERCIAL NF {dados['numero_nf']}.docx",
          dados,
          itens_jw,
          tot_jw,
          is_jw=True,
      )
      print(
          f"  ✓ JW criada (+{margem_jw*100:.2f}%) - Data:"
          f" {dados['data_emissao'].strftime('%d/%m/%Y')}"
      )


if __name__ == "__main__":
  processar_notas()